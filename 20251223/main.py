import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_document():
    doc = docx.Document()

    # 设置正文样式
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    # 中文字体设置通常需要复杂操作，这里主要确保西文正确，中文Word会自动回退到默认中文字体（如宋体/等线）

    # --- 标题 ---
    title = doc.add_heading('第二章 Transformer 模型结构分析', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 2.1 引言 ---
    doc.add_heading('2.1 引言', level=1)
    p = doc.add_paragraph(
        '在本章中，我们将深入分析 Vaswani 等人在论文《Attention Is All You Need》中提出的 Transformer 模型架构。'
        '作为一种全新的序列转导（Sequence Transduction）模型，Transformer 摒弃了传统的循环神经网络（RNN）和卷积神经网络（CNN）结构，'
        '完全基于注意力机制（Attention Mechanism）来处理序列依赖关系。这种设计不仅极大地提高了计算的并行度，'
        '还在机器翻译等任务上取得了当时的最优效果（SOTA）。本章将严格依据论文原文，从整体架构、注意力机制、前馈神经网络以及位置编码等方面进行详细阐述。'
    )

    # --- 2.2 整体架构 ---
    doc.add_heading('2.2 整体架构 (Model Architecture)', level=1)
    doc.add_paragraph(
        'Transformer 沿用了经典的编码器-解码器（Encoder-Decoder）结构。对于一个符号表示的输入序列 (x_1, ..., x_n)，'
        '编码器将其映射为连续表示序列 z = (z_1, ..., z_n)。给定 z，解码器则生成输出序列 (y_1, ..., y_m)。'
        '模型具有自回归（Auto-regressive）特性，即在生成每一个步骤的输出时，都会将之前生成的符号作为额外的输入。'
    )

    doc.add_heading('2.2.1 编码器 (Encoder)', level=2)
    doc.add_paragraph('编码器由 N=6 个相同的层堆叠而成。每一层包含两个子层：')
    doc.add_paragraph('1. 多头自注意力机制 (Multi-Head Self-Attention)：用于捕捉输入序列内部的依赖关系。',
                      style='List Number')
    doc.add_paragraph('2. 逐位置全连接前馈网络 (Position-wise Feed-Forward Networks)：对每个位置的向量进行非线性变换。',
                      style='List Number')
    doc.add_paragraph(
        '在这两个子层周围，Transformer 均引入了残差连接（Residual Connection），随后进行层归一化（Layer Normalization）。具体而言，每个子层的输出可以表示为：'
    )

    # 公式
    add_latex_formula(doc, r'\mathrm{LayerNorm}(x + \mathrm{Sublayer}(x))')

    doc.add_paragraph(
        '为了方便残差连接，模型中所有子层以及嵌入层（Embedding Layers）的输出维度均保持一致，即 d_{model} = 512。')

    doc.add_heading('2.2.2 解码器 (Decoder)', level=2)
    doc.add_paragraph('解码器同样由 N=6 个相同的层堆叠而成。与编码器不同的是，解码器的每一层包含三个子层：')
    doc.add_paragraph(
        '1. 掩码多头自注意力机制 (Masked Multi-Head Self-Attention)：该层通过掩码（Masking）机制，确保位置 i 的预测仅依赖于 i 之前的已知输出，从而保持自回归属性。',
        style='List Number')
    doc.add_paragraph(
        '2. 编码器-解码器注意力机制 (Encoder-Decoder Attention)：该层执行多头注意力计算，其中 Query 来自前一解码器层的输出，而 Key 和 Value 则来自编码器的输出。这使得解码器能够关注输入序列中的每一个位置。',
        style='List Number')
    doc.add_paragraph('3. 逐位置全连接前馈网络：与编码器中的结构相同。', style='List Number')
    doc.add_paragraph('同样，解码器的每个子层后也接有残差连接和层归一化。')

    # --- 2.3 注意力机制 ---
    doc.add_heading('2.3 注意力机制 (Attention Mechanism)', level=1)
    doc.add_paragraph(
        '注意力机制是 Transformer 的核心。论文将其描述为将一个 Query 和一组 Key-Value 对映射到一个输出的过程。')

    doc.add_heading('2.3.1 缩放点积注意力 (Scaled Dot-Product Attention)', level=2)
    doc.add_paragraph(
        '这是 Transformer 中最基本的注意力单元。输入包含维度为 d_k 的 Queries 和 Keys，以及维度为 d_v 的 Values。计算过程如下：'
    )
    doc.add_paragraph('1. 计算 Query 与所有 Keys 的点积。', style='List Number')
    doc.add_paragraph('2. 将结果除以 \\sqrt{d_k} 进行缩放。这是为了防止点积结果过大导致 Softmax 函数进入梯度极小的区域。',
                      style='List Number')
    doc.add_paragraph('3. 应用 Softmax 函数获取 Values 的权重。', style='List Number')

    doc.add_paragraph('矩阵形式的计算公式为：')
    add_latex_formula(doc, r'\mathrm{Attention}(Q, K, V) = \mathrm{softmax}\left(\frac{QK^T}{\sqrt{d_k}}\right)V')

    doc.add_heading('2.3.2 多头注意力 (Multi-Head Attention)', level=2)
    doc.add_paragraph(
        '为了让模型能够从不同的表示子空间（Representation Subspaces）关注不同位置的信息，论文提出了多头注意力机制。'
        '它不再使用单一的注意力函数，而是通过 h=8 个不同的线性投影（Linear Projections）将 Queries、Keys 和 Values 分别投影到 d_k、d_k 和 d_v 维度。'
    )
    add_latex_formula(doc, r'\mathrm{MultiHead}(Q, K, V) = \mathrm{Concat}(\mathrm{head}_1, ..., \mathrm{head}_h)W^O')
    doc.add_paragraph('其中，每个头的计算为：')
    add_latex_formula(doc, r'\mathrm{head}_i = \mathrm{Attention}(QW_i^Q, KW_i^K, VW_i^V)')
    doc.add_paragraph(
        '在论文的基础模型中，d_k = d_v = d_{model}/h = 64。这种设计使得多头注意力的总计算成本与全维度的单头注意力相当。')

    # --- 2.4 前馈网络 ---
    doc.add_heading('2.4 逐位置前馈网络 (Position-wise Feed-Forward Networks)', level=1)
    doc.add_paragraph(
        '除了注意力子层，编码器和解码器的每一层都包含一个全连接的前馈网络。该网络分别且独立地应用于每一个位置。它由两个线性变换和一个 ReLU 激活函数组成：'
    )
    add_latex_formula(doc, r'\mathrm{FFN}(x) = \max(0, xW_1 + b_1)W_2 + b_2')
    doc.add_paragraph(
        '虽然这个线性变换在不同位置上是相同的，但在不同的层之间采用了不同的参数。这可以被理解为核大小为 1 的两次卷积操作。'
        '输入和输出的维度为 d_{model} = 512，内部层的维度为 d_{ff} = 2048。'
    )

    # --- 2.5 位置编码 ---
    doc.add_heading('2.5 位置编码 (Positional Encoding)', level=1)
    doc.add_paragraph(
        '由于 Transformer 模型完全不包含循环（Recurrence）和卷积（Convolution），模型本身无法捕捉序列的顺序信息。'
        '为了解决这个问题，论文引入了“位置编码”，将其与输入嵌入（Input Embeddings）相加。'
    )
    doc.add_paragraph('位置编码具有与嵌入相同的维度 d_{model}。论文采用了正弦和余弦函数来生成位置编码：')
    add_latex_formula(doc, r'PE_{(pos, 2i)} = \sin(pos / 10000^{2i/d_{model}})')
    add_latex_formula(doc, r'PE_{(pos, 2i+1)} = \cos(pos / 10000^{2i/d_{model}})')
    doc.add_paragraph(
        '其中 pos 表示位置，i 表示维度。这种设计允许模型轻松学习相对位置的关注机制，因为对于任何固定的偏移量 k，'
        'PE_{pos+k} 都可以表示为 PE_{pos} 的线性函数。此外，这种正弦编码方式还有助于模型推断比训练期间遇到的序列更长的序列。'
    )

    # --- 2.6 总结 ---
    doc.add_heading('2.6 为什么选择自注意力 (Why Self-Attention)', level=1)
    doc.add_paragraph('论文从三个方面对比了自注意力层与循环层、卷积层：')
    doc.add_paragraph(
        '1. 每层的总计算复杂度：当序列长度 n 小于表示维度 d 时（这在现代机器翻译模型中很常见），自注意力层比循环层更快。',
        style='List Number')
    doc.add_paragraph('2. 可并行化的计算量：自注意力机制需要的顺序操作数为 O(1)，而循环神经网络为 O(n)。',
                      style='List Number')
    doc.add_paragraph(
        '3. 网络中长距离依赖之间的路径长度：学习长距离依赖的关键在于信号在网络中向前或向后传递的路径长度。在自注意力机制中，任意两个位置之间的最大路径长度仅为 O(1)，这使得模型能够更容易地学习长距离的依赖关系。',
        style='List Number')
    doc.add_paragraph(
        '综上所述，Transformer 通过独特的架构设计，在保证模型表达能力的同时，显著提升了训练效率和捕捉长距离依赖的能力。'
    )

    # 保存文档
    file_name = '课程论文_第二章_Transformer.docx'
    doc.save(file_name)
    print(f"文档已生成: {file_name}")


def add_latex_formula(doc, latex_text):
    """
    添加一个包含 LaTeX 代码的段落，格式化为居中，蓝色，方便用户识别并转换为公式。
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(latex_text)
    run.font.name = 'Cambria Math'  # 建议用于数学的字体
    run.font.color.rgb = RGBColor(0, 50, 150)  # 深蓝色以示区别
    run.italic = True


if __name__ == "__main__":
    create_document()